"""
Word (.docx) 문제지 자동 생성 모듈

출력 구성:
  1. 표지 (제목 / 출처 / 범위 / 날짜)
  2. 객관식 문제 파트
  3. 서술형 문제 파트 (선택)
  4. 답안지 (문항번호 + 정답)

의존 패키지:
  pip install python-docx
"""

import json
import re
from datetime import date
from typing import List, Dict, Any, Optional
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
    _LEFT   = WD_ALIGN_PARAGRAPH.LEFT
    _CENTER = WD_ALIGN_PARAGRAPH.CENTER
except ImportError:
    DOCX_AVAILABLE = False
    _LEFT = 0
    _CENTER = 1

from pipeline.config import OUTPUT_DIR


# ─── 헬퍼 ─────────────────────────────────────────────────────────────────────

CIRCLE_CHARS = ["①", "②", "③", "④", "⑤"]


def _get_choices(q: Dict[str, Any]) -> List[str]:
    choices = q.get("choices") or []
    if isinstance(choices, str):
        try:
            choices = json.loads(choices)
        except Exception:
            choices = []
    return choices


def _set_font(run, font_name: str = "나눔고딕", size_pt: int = 10):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)


def _add_paragraph(doc: "Document", text: str, bold: bool = False,
                   font_size: int = 10, alignment=_LEFT,
                   space_before: float = 0, space_after: float = 2) -> "Paragraph":
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)

    # <u> 태그가 있으면 _add_rich_runs로 처리
    if '<u>' in text:
        _add_rich_runs(p, text, font_size=font_size, bold=bold)
    else:
        run = p.add_run(text)
        run.bold = bold
        if DOCX_AVAILABLE:
            _set_font(run, size_pt=font_size)
    return p


def _sanitize_brackets(text: str) -> str:
    """<보기>, <조건> 등 한글 꺾쇠 라벨을 보기 좋은 괄호로 변환."""
    return re.sub(r'<(보기|조건|예시|참고|주의|힌트)>', r'[\1]', text)


def _add_rich_runs(paragraph, text: str, font_size: int = 9,
                   font_color: Optional["RGBColor"] = None,
                   bold: bool = False):
    """<u>...</u> 태그를 실제 밑줄 서식으로 변환하여 paragraph에 run을 추가한다.

    일반 텍스트는 그대로, <u>로 감싼 부분은 밑줄 서식을 적용한다.
    <보기> 등 한글 꺾쇠 라벨도 자동으로 변환한다.
    """
    text = _sanitize_brackets(text)

    # <u>...</u> 기준으로 분리  
    # 패턴: (<u>...</u>) 또는 일반 텍스트
    parts = re.split(r'(<u>.*?</u>)', text, flags=re.DOTALL)

    for part in parts:
        if not part:
            continue
        m = re.match(r'<u>(.*?)</u>', part, flags=re.DOTALL)
        if m:
            # 밑줄 텍스트
            run = paragraph.add_run(m.group(1))
            run.underline = True
            run.bold = bold
            if DOCX_AVAILABLE:
                _set_font(run, size_pt=font_size)
                if font_color:
                    run.font.color.rgb = font_color
        else:
            # 일반 텍스트
            run = paragraph.add_run(part)
            run.bold = bold
            if DOCX_AVAILABLE:
                _set_font(run, size_pt=font_size)
                if font_color:
                    run.font.color.rgb = font_color


# ─── 표지 ──────────────────────────────────────────────────────────────────────

def _add_cover(doc: "Document", title: str, info: str, today: str):
    doc.add_paragraph()
    _add_paragraph(doc, title, bold=True, font_size=18,
                   alignment=_CENTER, space_before=60, space_after=6)
    _add_paragraph(doc, info, font_size=11,
                   alignment=_CENTER, space_after=4)
    _add_paragraph(doc, today, font_size=10,
                   alignment=_CENTER, space_after=20)
    doc.add_page_break()


# ─── 객관식 문항 출력 ──────────────────────────────────────────────────────────

def _add_mc_question(doc: "Document", q: Dict[str, Any], display_no: int):
    """객관식 문항 하나를 문서에 추가."""
    choices = _get_choices(q)

    # 지문 — <u> 태그 밑줄 지원
    passage = q.get("passage_text", "")
    if passage:
        if '<u>' in passage:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            _add_rich_runs(p, passage, font_size=9)
        else:
            _add_paragraph(doc, passage, font_size=9, space_before=4, space_after=2)

    # 질문
    q_text = f"{display_no}. {q.get('question_text', '')}"
    _add_paragraph(doc, q_text, bold=False, font_size=10, space_before=4, space_after=2)

    # 선지
    for i, choice in enumerate(choices):
        if choice:
            _add_paragraph(doc, f"  {CIRCLE_CHARS[i]} {choice}", font_size=9.5, space_after=1)

    _add_paragraph(doc, "", space_after=4)  # 빈 줄


# ─── 서술형 문항 출력 ─────────────────────────────────────────────────────────

def _add_subjective_question(doc: "Document", q: Dict[str, Any], display_no: int):
    """서술형 문항 하나를 문서에 추가."""
    q_text = f"{display_no}. {q.get('question_text', '')}"
    _add_paragraph(doc, q_text, bold=False, font_size=10, space_before=6, space_after=2)
    # 답안 공간
    for _ in range(3):
        _add_paragraph(doc, "답: " + "_" * 60, font_size=10, space_after=2)
    _add_paragraph(doc, "", space_after=4)


# ─── 답안지 ────────────────────────────────────────────────────────────────────

def _add_answer_sheet(doc: "Document", questions: List[Dict[str, Any]]):
    doc.add_page_break()
    _add_paragraph(doc, "[ 정 답 ]", bold=True, font_size=13,
                   alignment=_CENTER, space_before=10, space_after=8)

    mc_qs = [q for q in questions if q.get("question_type") != "서술형"]
    subj_qs = [q for q in questions if q.get("question_type") == "서술형"]

    # 객관식 정답 표
    if mc_qs:
        tbl = doc.add_table(rows=2, cols=len(mc_qs))
        tbl.style = "Table Grid"
        # 헤더: 문항번호
        for i, q in enumerate(mc_qs):
            cell = tbl.rows[0].cells[i]
            cell.text = str(q.get("_display_no", i + 1))
            cell.paragraphs[0].alignment = _CENTER
        # 정답
        for i, q in enumerate(mc_qs):
            ans = q.get("answer", "")
            if isinstance(ans, str) and ans.isdigit():
                ans_display = CIRCLE_CHARS[int(ans) - 1]
            else:
                ans_display = str(ans)
            cell = tbl.rows[1].cells[i]
            cell.text = ans_display
            cell.paragraphs[0].alignment = _CENTER

    # 서술형 정답
    if subj_qs:
        _add_paragraph(doc, "\n[서술형 모범 답안]", bold=True, font_size=11, space_before=10)
        for q in subj_qs:
            no = q.get("_display_no", "")
            ans = q.get("answer") or "(정답 미입력)"
            _add_paragraph(doc, f"{no}. {ans}", font_size=10, space_after=4)


# ─── 메인 출력 함수 ────────────────────────────────────────────────────────────

def generate_exam_docx(
    questions: List[Dict[str, Any]],
    output_name: str = "문제지",
    title: str = "고1 영어 문제지",
    info: str = "",
    include_cover: bool = True,
    include_answers: bool = True,
    output_dir: Optional[Path] = None,
) -> Path:
    """
    문항 리스트로 Word 문제지(.docx) 생성.

    Args:
        questions    : search 모듈에서 반환된 문항 dict 목록
        output_name  : 파일명 (확장자 제외)
        title        : 표지 제목
        info         : 표지 부제 (출처/범위 등)
        include_cover: 표지 포함 여부
        include_answers: 답안지 페이지 포함 여부
        output_dir   : 저장 경로 (None 이면 config.OUTPUT_DIR)

    Returns:
        생성된 .docx 파일 경로
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx 가 설치되지 않았습니다. pip install python-docx 를 실행하세요.")

    out_dir = Path(output_dir) if output_dir else OUTPUT_DIR
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"{output_name}.docx"

    doc = Document()

    # 여백 설정
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(3.0)

    today_str = date.today().strftime("%Y년 %m월 %d일")

    if include_cover:
        _add_cover(doc, title, info, today_str)

    # 문항 분리
    mc_qs   = [q for q in questions if q.get("question_type") != "서술형"]
    subj_qs = [q for q in questions if q.get("question_type") == "서술형"]

    # 객관식 파트
    if mc_qs:
        _add_paragraph(doc, "[ 객관식 ]", bold=True, font_size=12, space_before=4, space_after=6)
        for i, q in enumerate(mc_qs, 1):
            q["_display_no"] = i
            _add_mc_question(doc, q, i)

    # 서술형 파트
    if subj_qs:
        if mc_qs:
            doc.add_page_break()
        _add_paragraph(doc, "[ 서술형 ]", bold=True, font_size=12, space_before=4, space_after=6)
        for i, q in enumerate(subj_qs, len(mc_qs) + 1):
            q["_display_no"] = i
            _add_subjective_question(doc, q, i)

    # 답안지
    if include_answers:
        all_qs = mc_qs + subj_qs
        _add_answer_sheet(doc, all_qs)

    doc.save(str(out_path))
    print(f"[Word] 문제지 저장 완료: {out_path}")
    return out_path


# ─── AI 생성 문제용 Word 출력 ─────────────────────────────────────────────────

def _add_ai_mc_question(doc: "Document", q: Dict[str, Any], display_no: int):
    """AI 생성 객관식 문항 하나를 문서에 추가."""
    # 질문 (번호 + 문제)
    q_text = _sanitize_brackets(f"{display_no}. {q.get('question_text', '')}")
    _add_paragraph(doc, q_text, bold=True, font_size=10, space_before=8, space_after=2)

    # 지문 — <u> 태그를 실제 밑줄 서식으로 변환
    passage = q.get("passage", "")
    if passage:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.5)
        _add_rich_runs(p, passage, font_size=9,
                       font_color=RGBColor(0x33, 0x33, 0x33))

    # 선지
    choices = q.get("choices", [])
    if isinstance(choices, str):
        try:
            choices = json.loads(choices)
        except Exception:
            choices = []
    for choice in choices:
        if choice:
            choice_text = _sanitize_brackets(f"    {choice}")
            if '<u>' in choice_text:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(1)
                _add_rich_runs(p, choice_text, font_size=9.5)
            else:
                _add_paragraph(doc, choice_text, font_size=9.5, space_after=1)

    _add_paragraph(doc, "", space_after=2)


def _add_ai_subj_question(doc: "Document", q: Dict[str, Any], display_no: int):
    """AI 생성 주관식 문항 하나를 문서에 추가."""
    q_text = _sanitize_brackets(f"{display_no}. {q.get('question_text', '')}")
    _add_paragraph(doc, q_text, bold=True, font_size=10, space_before=8, space_after=2)

    passage = q.get("passage", "")
    if passage:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.5)
        _add_rich_runs(p, passage, font_size=9,
                       font_color=RGBColor(0x33, 0x33, 0x33))

    # 답안 작성 공간
    _add_paragraph(doc, "", space_after=2)
    _add_paragraph(doc, "답: " + "_" * 60, font_size=10, space_after=6)


def _add_ai_answer_sheet(doc: "Document", questions: List[Dict[str, Any]],
                         question_type: str):
    """AI 생성 문항의 정답·해설 페이지."""
    doc.add_page_break()
    _add_paragraph(doc, "[ 정답 및 해설 ]", bold=True, font_size=14,
                   alignment=_CENTER, space_before=10, space_after=12)

    for q in questions:
        no = q.get("_display_no", q.get("question_no", ""))
        ans = q.get("answer", "")
        explanation = q.get("explanation", "")

        # 정답 표시
        if question_type == "객관식" and isinstance(ans, int) and 1 <= ans <= 5:
            ans_display = f"{CIRCLE_CHARS[ans - 1]}  ({ans}번)"
        else:
            ans_display = str(ans)

        _add_paragraph(doc, f"{no}번  정답: {ans_display}",
                       bold=True, font_size=10, space_before=6, space_after=1)

        # 해설 — <u> 태그 밑줄 지원
        if explanation:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Cm(0.8)
            _add_rich_runs(p, explanation, font_size=9,
                           font_color=RGBColor(0x55, 0x55, 0x55))

        # [해석] — <u> 태그 밑줄 지원
        translation = q.get("translation", "")
        if translation:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Cm(0.8)
            run_label = p.add_run("[해석] ")
            if DOCX_AVAILABLE:
                _set_font(run_label, size_pt=9)
                run_label.bold = True
                run_label.font.color.rgb = RGBColor(0x29, 0x80, 0xB9)
            _add_rich_runs(p, translation, font_size=9,
                           font_color=RGBColor(0x29, 0x80, 0xB9))


def generate_ai_exam_docx(
    result: Dict[str, Any],
    output_name: Optional[str] = None,
    include_answers: bool = True,
    output_dir: Optional[Path] = None,
) -> Path:
    """
    AI 생성 결과(generate_questions 반환값)를 Word 문제지로 출력.

    Args:
        result       : generate_questions()가 반환한 dict
        output_name  : 파일명 (확장자 제외, None이면 자동 생성)
        include_answers: 정답·해설 페이지 포함 여부
        output_dir   : 저장 경로 (None이면 config.OUTPUT_DIR)

    Returns:
        생성된 .docx 파일 경로
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx 가 설치되지 않았습니다. pip install python-docx")

    if not result.get("success"):
        raise ValueError(f"생성 실패 결과는 Word로 출력할 수 없습니다: {result.get('error')}")

    meta = result["metadata"]
    tb = result["textbook_used"]
    questions = result["questions"]
    question_type = meta["question_type"]

    out_dir = Path(output_dir) if output_dir else OUTPUT_DIR
    out_dir.mkdir(parents=True, exist_ok=True)

    # 파일명 자동 생성
    if not output_name:
        import datetime
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_name = (
            f"AI문제_{meta['sub_type']}_{question_type}_"
            f"{meta['count_generated']}문항_{ts}"
        )

    out_path = out_dir / f"{output_name}.docx"

    doc = Document()

    # 여백 설정
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── 표지 ──
    doc.add_paragraph()
    title = f"고1 영어 {meta['sub_type']} ({question_type})"
    _add_paragraph(doc, title, bold=True, font_size=18,
                   alignment=_CENTER, space_before=40, space_after=8)

    # 교과서 정보
    tb_label = tb.get('textbook_label', tb.get('label', ''))
    tb_unit = tb.get('unit_no', '')
    if tb_label and tb_unit:
        tb_info = f"{tb_label}  {tb_unit}과"
        if tb.get("unit_title"):
            tb_info += f" — {tb['unit_title']}"
    elif tb_label:
        tb_info = tb_label
    else:
        tb_info = tb.get('source', '모의고사 지문 활용')
    _add_paragraph(doc, tb_info, font_size=11,
                   alignment=_CENTER, space_after=4)

    # 부제
    sub_info_parts = [
        f"과목: {tb.get('subject', '')}",
        f"출판사: {tb.get('publisher', '')}",
        f"문항 수: {meta['count_generated']}",
    ]
    sub_info_parts = [p for p in sub_info_parts if not p.endswith(': ')]
    _add_paragraph(doc, "  |  ".join(sub_info_parts), font_size=9,
                   alignment=_CENTER, space_after=4)

    today_str = date.today().strftime("%Y년 %m월 %d일")
    _add_paragraph(doc, today_str, font_size=9,
                   alignment=_CENTER, space_after=16)

    # 이름 / 점수 란
    _add_paragraph(doc, "이름: _______________     점수: _____ / 100",
                   font_size=10, alignment=_CENTER, space_after=8)

    doc.add_page_break()

    # ── 문제 파트 ──
    section_label = "[ 객관식 ]" if question_type == "객관식" else "[ 서술형 ]"
    _add_paragraph(doc, section_label, bold=True, font_size=13,
                   space_before=4, space_after=8)

    for i, q in enumerate(questions, 1):
        q["_display_no"] = i
        if question_type == "객관식":
            _add_ai_mc_question(doc, q, i)
        else:
            _add_ai_subj_question(doc, q, i)

    # ── 정답·해설 ──
    if include_answers:
        _add_ai_answer_sheet(doc, questions, question_type)

    doc.save(str(out_path))
    return out_path


def generate_ai_exam_docx_multi(
    results: List[Dict[str, Any]],
    output_name: Optional[str] = None,
    include_answers: bool = True,
    output_dir: Optional[Path] = None,
) -> Path:
    """
    여러 유형의 AI 생성 결과를 하나의 Word 문제지로 합쳐서 출력.

    Args:
        results      : generate_questions() 반환값 리스트 (여러 유형)
        output_name  : 파일명 (확장자 제외, None이면 자동 생성)
        include_answers: 정답·해설 페이지 포함 여부
        output_dir   : 저장 경로 (None이면 config.OUTPUT_DIR)

    Returns:
        생성된 .docx 파일 경로
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx 가 설치되지 않았습니다. pip install python-docx")

    if not results:
        raise ValueError("생성 결과가 없습니다.")

    # 첫 번째 결과에서 교과서 정보 추출
    first_tb = results[0]["textbook_used"]
    first_meta = results[0]["metadata"]
    question_type = first_meta["question_type"]

    # 유형 목록, 총 문항 수 집계
    type_names = [r["metadata"]["sub_type"] for r in results]
    total_count = sum(r["metadata"]["count_generated"] for r in results)

    out_dir = Path(output_dir) if output_dir else OUTPUT_DIR
    out_dir.mkdir(parents=True, exist_ok=True)

    if not output_name:
        import datetime
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        types_str = "+".join(type_names)
        output_name = (
            f"AI문제_{types_str}_{question_type}_"
            f"{total_count}문항_{ts}"
        )

    out_path = out_dir / f"{output_name}.docx"

    doc = Document()

    # 여백
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── 표지 ──
    doc.add_paragraph()
    title = f"고1 영어 ({question_type})"
    _add_paragraph(doc, title, bold=True, font_size=18,
                   alignment=_CENTER, space_before=40, space_after=8)

    # 유형 목록
    types_display = " / ".join(type_names)
    _add_paragraph(doc, types_display, font_size=12,
                   alignment=_CENTER, space_after=6)

    # 교과서 정보
    tb_label = first_tb.get('textbook_label', first_tb.get('label', ''))
    tb_unit = first_tb.get('unit_no', '')
    if tb_label and tb_unit:
        tb_info = f"{tb_label}  {tb_unit}과"
        if first_tb.get("unit_title"):
            tb_info += f" — {first_tb['unit_title']}"
    elif tb_label:
        tb_info = tb_label
    else:
        tb_info = first_tb.get('source', '모의고사 지문 활용')
    _add_paragraph(doc, tb_info, font_size=11,
                   alignment=_CENTER, space_after=4)

    sub_info_parts = [
        f"과목: {first_tb.get('subject', '')}",
        f"출판사: {first_tb.get('publisher', '')}",
        f"총 {total_count}문항",
    ]
    sub_info_parts = [p for p in sub_info_parts if not p.endswith(': ')]
    _add_paragraph(doc, "  |  ".join(sub_info_parts), font_size=9,
                   alignment=_CENTER, space_after=4)

    today_str = date.today().strftime("%Y년 %m월 %d일")
    _add_paragraph(doc, today_str, font_size=9,
                   alignment=_CENTER, space_after=16)

    _add_paragraph(doc, "이름: _______________     점수: _____ / 100",
                   font_size=10, alignment=_CENTER, space_after=8)

    doc.add_page_break()

    # ── 문제 파트 (유형별 섹션) ──
    all_questions = []
    global_no = 1

    for result in results:
        meta = result["metadata"]
        questions = result["questions"]
        sub_type = meta["sub_type"]
        source_label = meta.get("source_label", "")

        # 유형 헤더 (소스 정보 포함)
        section_label = f"[ {sub_type} ]"
        if source_label and source_label != "전체":
            section_label += f"  — {source_label} 기반"
        if question_type == "주관식":
            section_label += " (서술형)"
        _add_paragraph(doc, section_label, bold=True, font_size=13,
                       space_before=10, space_after=8)

        for q in questions:
            q["_display_no"] = global_no
            if question_type == "객관식":
                _add_ai_mc_question(doc, q, global_no)
            else:
                _add_ai_subj_question(doc, q, global_no)
            all_questions.append(q)
            global_no += 1

    # ── 정답·해설 ──
    if include_answers:
        _add_ai_answer_sheet(doc, all_questions, question_type)

    doc.save(str(out_path))
    return out_path
